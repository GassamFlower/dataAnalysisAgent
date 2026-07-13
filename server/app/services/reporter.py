"""报告导出服务。

Word（python-docx）+ Excel（openpyxl）。
所有导出文件强制带 simulated 水印 + 免责声明。
分档标签与论文段落模板来源：app/core/statistics_constants.py
"""
import io
from datetime import datetime
from typing import Dict, List, Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill

from app.core.statistics_constants import (
    GRADE_TABLE_TEXT,
    grade_alpha,
    grade_bartlett,
    grade_kmo,
)


def _to_float(value: Any, default: float = 0.0) -> float:
    """安全转 float（DB 中可能为 Decimal/str）。"""
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _format_p_value(p: Any) -> str:
    """格式化 p 值：<0.001 显示为 p<0.001，否则保留 3 位小数。"""
    try:
        p_val = float(p)
    except (TypeError, ValueError):
        return "—"
    if p_val < 0.001:
        return "<0.001"
    return f"{p_val:.3f}"


def _format_statistic(stat: Any) -> str:
    """格式化统计量：保留 3 位小数，无效值显示「—」。"""
    try:
        return f"{float(stat):.3f}"
    except (TypeError, ValueError):
        return "—"


def _format_effect_size(value: Any, name: str, grade: str) -> str:
    """格式化效应量：值（名称，分档）。"""
    try:
        val = float(value)
        return f"{val:.3f}（{name or '—'}，{grade or '—'}）"
    except (TypeError, ValueError):
        return "—"


def _diff_test_table_to_doc(doc: "Document", diff_tests: List[Dict[str, Any]]) -> None:
    """在 Word 文档中插入差异检验结果章节。

    含表格 + 每条路径的自然语言解读。
    设计依据：docs/后端架构设计文档.md 第 9.6 节决策树。
    """
    if not diff_tests:
        doc.add_paragraph(
            "未配置假设路径，无差异检验结果。请在「生成数据」步骤中填写研究假设。"
        )
        return

    # 差异检验结果表
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "假设路径"
    hdr_cells[1].text = "检验方法"
    hdr_cells[2].text = "统计量"
    hdr_cells[3].text = "p 值"
    hdr_cells[4].text = "效应量"
    hdr_cells[5].text = "显著性"

    for t in diff_tests:
        predictor = t.get("predictor", "")
        outcome = t.get("outcome", "")
        path_label = f"{predictor} → {outcome}"
        method_name = t.get("method_name") or t.get("method") or "—"

        # 错误场景：只显示原因
        if t.get("error"):
            row_cells = table.add_row().cells
            row_cells[0].text = path_label
            row_cells[1].text = method_name
            for i in range(2, 6):
                row_cells[i].text = "—"
            row_cells[3].text = f"错误：{t['error']}"
            continue

        significant = t.get("significant")
        sig_text = "显著 *" if significant else "不显著"

        row_cells = table.add_row().cells
        row_cells[0].text = path_label
        row_cells[1].text = method_name
        row_cells[2].text = _format_statistic(t.get("statistic"))
        row_cells[3].text = _format_p_value(t.get("p_value"))
        row_cells[4].text = _format_effect_size(
            t.get("effect_size"), t.get("effect_size_name"), t.get("effect_size_grade")
        )
        row_cells[5].text = sig_text

    # 自然语言解读区
    doc.add_paragraph()
    doc.add_heading("差异检验结果解读", level=2)
    has_interp = False
    for t in diff_tests:
        interpretation = t.get("interpretation")
        if not interpretation:
            continue
        has_interp = True
        predictor = t.get("predictor", "")
        outcome = t.get("outcome", "")
        p = doc.add_paragraph()
        p.add_run(f"{predictor} → {outcome}：").bold = True
        p.add_run(interpretation)

    if not has_interp:
        doc.add_paragraph("无可用的自然语言解读。")


def _reliability_paragraph(report_data: Dict[str, Any]) -> str:
    """套用论文信效度段落模板（一段式），填入计算结果。

    缺失字段以「—」占位，保证段落结构完整可读。
    设计依据：docs/后端架构设计文档.md 第 9.4 节
    """
    reliability = report_data.get("reliability_results", []) or []
    dim_count = len(reliability)
    overall_alpha = _to_float(report_data.get("overall_alpha", 0))

    if reliability:
        alphas = [_to_float(r.get("alpha", 0)) for r in reliability]
        kmos = [_to_float(r.get("kmo", 0)) for r in reliability]
        bartletts = [_to_float(r.get("bartlett_p_value", 1)) for r in reliability]
        min_alpha = min(alphas)
        max_alpha = max(alphas)
        avg_kmo = sum(kmos) / len(kmos) if kmos else 0
        bartlett_pass = all(b < 0.05 for b in bartletts)
        _, alpha_wording = grade_alpha(overall_alpha)
    else:
        min_alpha = max_alpha = avg_kmo = 0
        bartlett_pass = False
        alpha_wording = "无数据"

    # 因子数与累计方差在当前流程未持久化，按维度数推断并占位
    factor_count = dim_count if dim_count else "—"
    variance_pct = "—"

    paragraph = (
        f"本量表共 {dim_count or '—'} 个维度。"
        f"信度检验显示，总量表 Cronbach's α = {overall_alpha:.3f}（{alpha_wording}），"
        f"各维度 α 介于 {min_alpha:.3f}～{max_alpha:.3f}。"
        f"效度检验中，KMO = {avg_kmo:.3f}，"
        f"Bartlett 球形检验 p{'<0.05' if bartlett_pass else '≥0.05'}，"
        f"{'适合做因子分析' if bartlett_pass and avg_kmo >= 0.5 else '因子分析适用性需进一步评估'}；"
        f"主成分分析提取 {factor_count} 个公因子，累计方差解释率 {variance_pct}。"
    )
    return paragraph


def export_word(report_data: Dict[str, Any]) -> bytes:
    """导出 Word 报告。

    含 simulated 水印、免责声明、统计结果、R4 诊断结论。
    """
    doc = Document()

    # 添加水印
    _add_watermark(doc)

    # 标题
    title = doc.add_heading("数据分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 报告元信息
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"项目 ID：{report_data.get('project_id', 'N/A')}")
    doc.add_paragraph()

    # 总体结果
    doc.add_heading("一、总体信效度", level=1)
    overall_alpha = report_data.get("overall_alpha", 0)
    passed_count = report_data.get("passed_count", 0)
    total_count = report_data.get("total_count", 0)

    doc.add_paragraph(f"整体 Cronbach's α：{overall_alpha}")
    doc.add_paragraph(f"通过维度数：{passed_count} / {total_count}")
    doc.add_paragraph()

    # 各维度信效度结果
    doc.add_heading("二、各维度信效度详情", level=1)
    reliability_results = report_data.get("reliability_results", [])

    if reliability_results:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "维度"
        hdr_cells[1].text = "Cronbach's α"
        hdr_cells[2].text = "KMO"
        hdr_cells[3].text = "Bartlett p"
        hdr_cells[4].text = "是否通过"

        for r in reliability_results:
            row_cells = table.add_row().cells
            row_cells[0].text = str(r.get("dimension", ""))
            # 数值后附加分档等级（来源：statistics_constants）
            alpha_val = _to_float(r.get("alpha"))
            kmo_val = _to_float(r.get("kmo"))
            bartlett_val = _to_float(r.get("bartlett_p_value"))
            alpha_grade, _ = grade_alpha(alpha_val)
            kmo_grade, _ = grade_kmo(kmo_val)
            bartlett_grade, _ = grade_bartlett(bartlett_val)
            row_cells[1].text = f"{alpha_val:.3f} ({alpha_grade})"
            row_cells[2].text = f"{kmo_val:.3f} ({kmo_grade})"
            row_cells[3].text = f"{bartlett_val:.5f} ({bartlett_grade})"
            row_cells[4].text = "✓" if r.get("passed") else "✗"
    else:
        doc.add_paragraph("无信效度数据")

    doc.add_paragraph()

    # 论文信效度段落（参考，可直接进论文方法部分）
    doc.add_heading("三、论文信效度段落（参考）", level=1)
    doc.add_paragraph(_reliability_paragraph(report_data))
    doc.add_paragraph()

    # 差异检验（假设路径验证，对应架构文档 9.6 决策树）
    doc.add_heading("四、假设检验（差异分析）", level=1)
    doc.add_paragraph(
        "按假设路径自动选择检验方法（t检验/ANOVA/卡方/Pearson/线性回归），"
        "结果实时计算，不落库。"
    )
    _diff_test_table_to_doc(doc, report_data.get("diff_tests", []) or [])
    doc.add_paragraph()

    # R4 诊断结论
    doc.add_heading("五、R4 诊断结论", level=1)
    diagnosis = report_data.get("diagnosis")

    if diagnosis:
        passed = diagnosis.get("passed", False)
        status_text = "通过" if passed else "不通过"
        doc.add_paragraph(f"诊断结果：{status_text}")

        issues = diagnosis.get("issues", [])
        if issues:
            doc.add_paragraph(f"发现问题数：{len(issues)}")
            doc.add_paragraph()

            for i, issue in enumerate(issues, 1):
                doc.add_heading(f"问题 {i}：{issue.get('dimension', '')} - {issue.get('metric', '')}", level=2)
                doc.add_paragraph(f"指标值：{issue.get('value', '')}")
                doc.add_paragraph(f"阈值：{issue.get('threshold', '')}")
                doc.add_paragraph(f"原因：{issue.get('reason', '')}")
                doc.add_paragraph(f"建议：{issue.get('suggestion', '')}")
                doc.add_paragraph()
        else:
            doc.add_paragraph("未发现显著问题。")
    else:
        doc.add_paragraph("无诊断数据")

    # 附录：信效度速查表
    doc.add_paragraph()
    doc.add_heading("附录：信效度速查表", level=1)
    appendix = doc.add_paragraph(GRADE_TABLE_TEXT)
    appendix.runs[0].font.size = Pt(9)
    appendix.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    # 免责声明
    doc.add_paragraph()
    doc.add_heading("免责声明", level=1)
    disclaimer = doc.add_paragraph(
        "本报告基于模拟数据生成，仅供学术研究和教学演示使用，不代表真实统计分析结果。"
        "实际数据分析请联系专业统计人员。"
    )
    disclaimer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    disclaimer.runs[0].font.size = Pt(9)

    # 保存到字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_excel(dataset: Dict[str, Any]) -> bytes:
    """导出 Excel 数据集。

    含 simulated 水印元数据 sheet。
    """
    wb = Workbook()

    # 元数据 sheet（含水印说明）
    ws_meta = wb.active
    ws_meta.title = "元数据"
    ws_meta["A1"] = "数据分析报告 - 模拟数据集"
    ws_meta["A2"] = f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ws_meta["A3"] = "注意：本数据为模拟数据，仅供演示使用"
    ws_meta["A4"] = f"项目 ID：{dataset.get('project_id', 'N/A')}"

    # 信效度结果 sheet
    ws_reliability = wb.create_sheet("信效度结果")
    reliability_results = dataset.get("reliability_results", [])

    if reliability_results:
        headers = ["维度", "Cronbach's α", "KMO", "Bartlett p", "是否通过"]
        ws_reliability.append(headers)

        # 表头样式
        for cell in ws_reliability[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")

        for r in reliability_results:
            ws_reliability.append([
                r.get("dimension", ""),
                float(r.get("alpha", 0)),
                float(r.get("kmo", 0)),
                float(r.get("bartlett_p_value", 0)),
                "通过" if r.get("passed") else "不通过"
            ])

        # 调整列宽
        ws_reliability.column_dimensions["A"].width = 20
        ws_reliability.column_dimensions["B"].width = 15
        ws_reliability.column_dimensions["C"].width = 10
        ws_reliability.column_dimensions["D"].width = 15
        ws_reliability.column_dimensions["E"].width = 12
    else:
        ws_reliability.append(["无信效度数据"])

    # 诊断问题 sheet
    ws_issues = wb.create_sheet("诊断问题")
    diagnosis = dataset.get("diagnosis")

    if diagnosis:
        issues = diagnosis.get("issues", [])
        if issues:
            headers = ["维度", "指标", "指标值", "阈值", "原因", "建议"]
            ws_issues.append(headers)

            # 表头样式
            for cell in ws_issues[1]:
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")

            for issue in issues:
                ws_issues.append([
                    issue.get("dimension", ""),
                    issue.get("metric", ""),
                    float(issue.get("value", 0)),
                    float(issue.get("threshold", 0)),
                    issue.get("reason", ""),
                    issue.get("suggestion", "")
                ])

            # 调整列宽
            ws_issues.column_dimensions["A"].width = 15
            ws_issues.column_dimensions["B"].width = 12
            ws_issues.column_dimensions["C"].width = 10
            ws_issues.column_dimensions["D"].width = 10
            ws_issues.column_dimensions["E"].width = 40
            ws_issues.column_dimensions["F"].width = 40
        else:
            ws_issues.append(["未发现显著问题"])
    else:
        ws_issues.append(["无诊断数据"])

    # 差异检验结果 sheet（对应架构文档 9.6 决策树）
    ws_diff = wb.create_sheet("差异检验")
    diff_tests = dataset.get("diff_tests", []) or []

    if diff_tests:
        headers = [
            "假设路径", "检验方法", "统计量", "p 值",
            "效应量", "效应量名称", "分档", "显著性", "解读",
        ]
        ws_diff.append(headers)

        # 表头样式
        for cell in ws_diff[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")

        for t in diff_tests:
            predictor = t.get("predictor", "")
            outcome = t.get("outcome", "")
            path_label = f"{predictor} → {outcome}"
            method_name = t.get("method_name") or t.get("method") or ""

            # 错误场景：仅写错误原因
            if t.get("error"):
                ws_diff.append([
                    path_label, method_name, "", "", "", "", "", "",
                    f"错误：{t['error']}",
                ])
                continue

            # 数值字段安全转 float
            try:
                stat_val = float(t.get("statistic")) if t.get("statistic") is not None else None
            except (TypeError, ValueError):
                stat_val = None
            try:
                p_val = float(t.get("p_value")) if t.get("p_value") is not None else None
            except (TypeError, ValueError):
                p_val = None
            try:
                es_val = float(t.get("effect_size")) if t.get("effect_size") is not None else None
            except (TypeError, ValueError):
                es_val = None

            significant = t.get("significant")
            sig_text = "显著 *" if significant is True else ("不显著" if significant is False else "")

            ws_diff.append([
                path_label,
                method_name,
                stat_val if stat_val is not None else "",
                p_val if p_val is not None else "",
                es_val if es_val is not None else "",
                t.get("effect_size_name", "") or "",
                t.get("effect_size_grade", "") or "",
                sig_text,
                t.get("interpretation", "") or "",
            ])

        # 调整列宽
        ws_diff.column_dimensions["A"].width = 22
        ws_diff.column_dimensions["B"].width = 14
        ws_diff.column_dimensions["C"].width = 12
        ws_diff.column_dimensions["D"].width = 12
        ws_diff.column_dimensions["E"].width = 12
        ws_diff.column_dimensions["F"].width = 14
        ws_diff.column_dimensions["G"].width = 10
        ws_diff.column_dimensions["H"].width = 10
        ws_diff.column_dimensions["I"].width = 50
    else:
        ws_diff.append(["未配置假设路径，无差异检验结果"])

    # 保存到字节流
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_dataset_excel(
    columns: List[str], data: List[Dict[str, Any]], meta: Dict[str, Any]
) -> bytes:
    """导出模拟数据集 Excel。

    Sheet1：模拟数据（原始数据行）
    Sheet2：元数据（simulated 水印 + 免责声明）
    """
    wb = Workbook()

    # Sheet1: 模拟数据
    ws_data = wb.active
    ws_data.title = "模拟数据"
    ws_data.append(columns)

    # 表头样式
    for cell in ws_data[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(
            start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
        )

    # 数据行
    for row in data:
        ws_data.append([row.get(col) for col in columns])

    # Sheet2: 元数据（含水印）
    ws_meta = wb.create_sheet("元数据")
    ws_meta["A1"] = "SIMULATED DATA - 模拟数据集"
    ws_meta["A1"].font = Font(bold=True, color="FF0000")
    ws_meta["A2"] = f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ws_meta["A3"] = f"项目 ID：{meta.get('project_id', 'N/A')}"
    ws_meta["A4"] = f"样本量：{len(data)}"
    ws_meta["A5"] = f"维度数：{len(columns)}"
    ws_meta["A7"] = "免责声明"
    ws_meta["A7"].font = Font(bold=True)
    ws_meta["A8"] = "本数据为模拟数据，仅供学术研究和教学演示使用，不代表真实数据。"
    ws_meta["A9"] = "实际数据采集请联系专业调研人员。"

    # 保存到字节流
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _add_watermark(doc: Document) -> None:
    """添加 simulated 水印（通过页眉实现）。"""
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "SIMULATED DATA - 仅供演示"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置水印样式
    run = paragraph.runs[0]
    run.font.color.rgb = RGBColor(192, 192, 192)
    run.font.size = Pt(10)
    run.font.italic = True
