"""问卷星导出文件解析服务单元测试。

覆盖（testing-strategy 第二步：三类用例）：
- 正常：Excel/Word 标准文件解析、题目提取、维度提取、题型映射
- 边界：空选项、缺失维度列、无标签题型、多选项行
- 异常：空文件、不支持的类型、损坏的文件内容、缺少依赖（word）

设计依据：docs/w-功能-问卷星链接解析.md
"""
import io

import pytest

from app.services.wjx_parser import (
    parse_wjx_export,
    _map_question_type,
    _parse_options,
    _detect_wjx_columns,
)


# ============================
# 辅助函数：构造测试文件
# ============================

def _build_wjx_excel_content() -> bytes:
    """构造问卷星导出的 Excel 文件内容（标准格式）。

    格式：
    | 题号 | 题目       | 题型   | 选项     | 维度     |
    | Q1   | 您的性别   | 单选   | 男,女    | 人口学   |
    | Q2   | 学习动机1  | 量表   | 1-5分    | 学习动机 |
    | Q3   | 请填写建议 | 填空   |          |         |
    """
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["题号", "题目", "题型", "选项", "维度"])
    ws.append(["Q1", "您的性别", "单选", "男,女", "人口学"])
    ws.append(["Q2", "学习动机1", "量表", "1-5分", "学习动机"])
    ws.append(["Q3", "请填写建议", "填空", None, None])

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _build_wjx_word_content() -> bytes:
    """构造问卷星导出的 Word 文件内容（标准格式）。

    格式：
    1. 您的性别 [单选题]
       ○ 男
       ○ 女

    2. 学习动机强度 [量表题]
       ○ 1 ○ 2 ○ 3 ○ 4 ○ 5
    """
    import docx

    doc = docx.Document()
    doc.add_paragraph("1. 您的性别 [单选题]")
    doc.add_paragraph("○ 男")
    doc.add_paragraph("○ 女")
    doc.add_paragraph("")
    doc.add_paragraph("2. 学习动机强度 [量表题]")
    doc.add_paragraph("○ 1 ○ 2 ○ 3 ○ 4 ○ 5")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================
# 正常用例 — Excel
# ============================

def test_parse_wjx_excel_standard_file():
    """正常：标准 Excel 文件解析成功，提取题目、选项和维度。"""
    content = _build_wjx_excel_content()
    result = parse_wjx_export(content, "excel")

    assert len(result["questions"]) == 3
    assert result["warnings"] == []

    # 验证第一题
    q1 = result["questions"][0]
    assert q1["number"] == "Q1"
    assert q1["text"] == "您的性别"
    assert q1["type"] == "single_choice"
    assert q1["options"] == ["男", "女"]
    assert q1["dimension"] == "人口学"
    assert q1["is_reverse"] is False

    # 验证第二题
    q2 = result["questions"][1]
    assert q2["type"] == "likert"
    assert q2["dimension"] == "学习动机"

    # 验证维度去重
    assert set(result["dimensions"]) == {"人口学", "学习动机"}


def test_parse_wjx_excel_question_type_mapping():
    """正常：题型映射正确（单选/量表/填空/多选/矩阵）。"""
    content = _build_wjx_excel_content()
    result = parse_wjx_export(content, "excel")

    types = [q["type"] for q in result["questions"]]
    assert types == ["single_choice", "likert", "open_ended"]


# ============================
# 正常用例 — Word
# ============================

def test_parse_wjx_word_standard_file():
    """正常：标准 Word 文件解析成功，提取题目和选项。"""
    content = _build_wjx_word_content()
    result = parse_wjx_export(content, "word")

    assert len(result["questions"]) == 2

    q1 = result["questions"][0]
    assert q1["index"] == 1
    assert q1["text"] == "您的性别"
    assert q1["type"] == "single_choice"
    assert q1["options"] == ["男", "女"]

    q2 = result["questions"][1]
    assert q2["text"] == "学习动机强度"
    assert q2["type"] == "likert"
    assert q2["options"] == ["1", "2", "3", "4", "5"]


def test_parse_wjx_word_multi_option_line():
    """正常：单行多选项（○ 1 ○ 2 ○ 3）正确拆分。"""
    content = _build_wjx_word_content()
    result = parse_wjx_export(content, "word")

    q2 = result["questions"][1]
    assert len(q2["options"]) == 5
    assert q2["options"][0] == "1"
    assert q2["options"][4] == "5"


# ============================
# 边界用例
# ============================

def test_parse_wjx_excel_empty_options():
    """边界：填空题无选项，options 应为空列表。"""
    content = _build_wjx_excel_content()
    result = parse_wjx_export(content, "excel")

    q3 = result["questions"][2]
    assert q3["options"] == []
    assert q3["dimension"] is None


def test_parse_wjx_excel_no_dimension_column():
    """边界：无维度列时，dimension 应为 None。"""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["题号", "题目", "题型", "选项"])
    ws.append(["Q1", "测试题", "单选", "A,B"])

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    result = parse_wjx_export(buffer.getvalue(), "excel")
    q1 = result["questions"][0]
    assert q1["dimension"] is None
    assert result["dimensions"] == []


def test_parse_wjx_excel_unknown_question_type():
    """边界：未知题型映射为 'unknown'。"""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["题号", "题目", "题型", "选项", "维度"])
    ws.append(["Q1", "新题型", "排序题", "A,B", "测试"])

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    result = parse_wjx_export(buffer.getvalue(), "excel")
    assert result["questions"][0]["type"] == "unknown"


def test_map_question_type_all_variants():
    """边界：题型映射覆盖所有变体（全称/简称）。"""
    assert _map_question_type("单选题") == "single_choice"
    assert _map_question_type("单选") == "single_choice"
    assert _map_question_type("多选题") == "multiple_choice"
    assert _map_question_type("多选") == "multiple_choice"
    assert _map_question_type("量表题") == "likert"
    assert _map_question_type("量表") == "likert"
    assert _map_question_type("矩阵题") == "matrix"
    assert _map_question_type("矩阵") == "matrix"
    assert _map_question_type("填空题") == "open_ended"
    assert _map_question_type("填空") == "open_ended"
    assert _map_question_type("未知类型") == "unknown"


def test_parse_options_newline_separated():
    """边界：换行分隔的选项正确解析。"""
    assert _parse_options("男\n女\n其他") == ["男", "女", "其他"]


def test_parse_options_empty():
    """边界：空选项返回空列表。"""
    assert _parse_options("") == []
    assert _parse_options(None) == []


def test_detect_wjx_columns_synonyms():
    """边界：列名同义词检测（编号/内容/类型/答案/分类）。"""
    import pandas as pd

    df = pd.DataFrame(columns=["编号", "内容", "类型", "答案", "分类"])
    mapping = _detect_wjx_columns(df)

    assert mapping["编号"] == "question_number"
    assert mapping["内容"] == "question_text"
    assert mapping["类型"] == "question_type"
    assert mapping["答案"] == "options"
    assert mapping["分类"] == "dimension"


# ============================
# 异常用例
# ============================

def test_parse_wjx_unsupported_file_type():
    """异常：不支持的文件类型应抛出 ValueError。"""
    with pytest.raises(ValueError, match="不支持的文件类型"):
        parse_wjx_export(b"fake content", "csv")


def test_parse_wjx_excel_empty_content():
    """异常：空 Excel 文件应抛出 ValueError。"""
    with pytest.raises(ValueError):
        parse_wjx_export(b"", "excel")


def test_parse_wjx_excel_corrupted_content():
    """异常：损坏的 Excel 内容应抛出 ValueError。"""
    with pytest.raises(ValueError, match="Excel 文件解析失败"):
        parse_wjx_export(b"not an excel file at all", "excel")


def test_parse_wjx_word_empty_content():
    """异常：空 Word 文件应抛出 ValueError。"""
    with pytest.raises(ValueError):
        parse_wjx_export(b"", "word")


def test_parse_wjx_excel_header_only_no_data():
    """异常：Excel 只有表头无数据行应抛出 ValueError（内容为空）。"""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    # 只有表头，无数据行
    ws.append(["题号", "题目", "题型", "选项", "维度"])

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    # pandas 读取只有表头的文件得到空 DataFrame，应报"内容为空"
    with pytest.raises(ValueError, match="内容为空"):
        parse_wjx_export(buffer.getvalue(), "excel")


def test_parse_wjx_word_no_questions():
    """异常：Word 无题目行（不匹配题号格式）应抛出 ValueError。"""
    import docx

    doc = docx.Document()
    doc.add_paragraph("这是一段没有题号格式的文字")
    doc.add_paragraph("另一段普通文字")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    with pytest.raises(ValueError, match="未能解析出任何题目"):
        parse_wjx_export(buffer.getvalue(), "word")
